import os
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
from docx import Document

def browse_folder():
    folder_selected = filedialog.askdirectory()
    if folder_selected:
        folder_path.set(folder_selected)

def start_scan():
    folder = folder_path.get()
    if not folder:
        messagebox.showerror("Error", "Please select a folder first.")
        return

    extension = simpledialog.askstring("File type", "Enter the file extension (e.g., py, cpp, java):")
    if not extension:
        messagebox.showerror("Error", "No extension provided.")
        return

    extension = extension.strip().lstrip(".")
    document = Document()
    document.add_heading(f"Collected files (*.{extension}) from {folder}", level=1)

    count = 0

    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.endswith(f".{extension}"):
                count += 1
                file_path = os.path.join(root, file)
                document.add_heading(file, level=2)
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        content = f.read()
                except Exception as e:
                    content = f"Could not read file due to error: {e}"

                document.add_paragraph(content)

    if count == 0:
        messagebox.showinfo("Done", f"No files with extension .{extension} found.")
        return

    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if save_path:
        document.save(save_path)
        messagebox.showinfo("Done", f"Successfully saved {count} files to {save_path}")

# GUI
root = tk.Tk()
root.title("Code Collector")

folder_path = tk.StringVar()

frame = tk.Frame(root, padx=20, pady=20)
frame.pack()

label = tk.Label(frame, text="Select folder to scan:")
label.pack()

entry = tk.Entry(frame, textvariable=folder_path, width=50)
entry.pack(side=tk.LEFT)

browse_button = tk.Button(frame, text="Browse", command=browse_folder)
browse_button.pack(side=tk.LEFT, padx=5)

start_button = tk.Button(root, text="Start Scan and Save to Word", command=start_scan, bg="lightblue")
start_button.pack(pady=20)

root.mainloop()
